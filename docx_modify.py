'''
(1)本程序先把doc转换为docx,再替换docx 文件里的特定字符串，包括正文、页眉、页脚、表格内容；
(2)空文件会报错；
(3)c:\目录下必须以管理员权限运行；

'''

import os
from docx import Document
import time
from win32com import client as wc




old_name = [r'机密',r'秘密',r'绝密']  #待替换关键字
new_name = r'内部使用'              #新关键字
path = r'C:\Users\Administrator\Desktop\py'
err_log =  u'log列表.txt'

def log(text):
    with open( err_log,"a+" ,encoding="utf-8") as f:
        print(text)
        try:
            f.write(text)
        except Exception as e:
            f.write("%s"%(repr(e)))
        f.write('\n')
    
change_count = 0
def replace_str(keys,target_str,fileName):
    global change_count
    str_new = None
    for key in keys:
        if (key in target_str):
            str_new = target_str.replace(key, new_name)
            target_str = str_new
            timeStr = time.strftime("%Y-%M-%D_%H%M%S",time.localtime())
            change_count = change_count+1
            log("%s %s ---内容修改%d,%s"%(timeStr,fileName,change_count,key))  
    
    return str_new

def doc_to_docx(path):
    w = wc.Dispatch('Word.Application')
    # 或者使用下面的方法，使用启动独立的进程：
    # w = wc.DispatchEx('Word.Application')
    for parent, dirnames, filenames in os.walk(path):
        for fn in filenames:
            filedir = os.path.join(parent, fn)
            if fn.endswith('.doc'):
                try:
                    doc=w.Documents.Open(filedir)
                    new_file = filedir.replace(r".doc",r".docx")
                    doc.SaveAs(new_file,16)#必须有参数16，否则会出错
                    doc.Close() #关闭原来word文件
                    os.remove(filedir)
                    log("%s => %s ---doc转换docx"%(filedir,new_file))
                except Exception as e:
                    log("%s %s ---转换错误，忽略"%(filedir,repr(e)))
                    continue
                
    w.Quit()
        
def change_header(path):
    for parent, dirnames, filenames in os.walk(path):
        for fn in filenames:
            filedir = os.path.join(parent, fn)
            if fn.endswith('.docx'):
                log("%s ---遍历文件"%(filedir))
                if(0 == os.path.getsize(filedir)): #忽略空文件
                    continue
                try:
                    document = Document(filedir) #打开文档
                except Exception as e:
                    log("%s ---打开错误，忽略"%(repr(e)))
                    continue

                try:
                    for p in document.paragraphs:  #正文
                        newText=replace_str(old_name,p.text,filedir)
                        if(newText != None):
                            p.text = newText

                    for table in document.tables:   #表格
                            for row in table.rows:
                                for cell in row.cells:
                                    newText=replace_str(old_name,cell.text,filedir)
                                    if(newText != None):
                                        cell.text = newText
                    
                    for section in document.sections:
                        header = section.header # 获取页眉
                        #print('页眉中默认段落数：', len(header.paragraphs))
                        for paragraph in header.paragraphs:          
                            newText=replace_str(old_name,paragraph.text,filedir)
                            if(newText != None):
                                paragraph.text = newText
                        
                        footer = section.footer # 获取页脚
                        #print('页脚中默认段落数：', len(footer.paragraphs))
                        for paragraph in footer.paragraphs:          
                            newText=replace_str(old_name,paragraph.text,filedir)
                            if(newText != None):
                                paragraph.text = newText
                except Exception as e:
                    log("%s ---编辑错误，忽略"%(repr(e)))
                    
                document.save(filedir) # 保存文档

if __name__ == "__main__":
    print("please input dir to walk,default is current dir:")
    path = input()
    if(0 == len(path)):
        path = os.getcwd()
    doc_to_docx(path)
    change_header(path)
    print("docx modify done,modify count=%d.press any key to exit."%(change_count))
    input()
    

